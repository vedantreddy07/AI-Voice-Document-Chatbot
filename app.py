import streamlit as st
import numpy as np
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re
import datetime
import tempfile
import os
from io import BytesIO
import PyPDF2
import docx
import json
from typing import List, Dict, Tuple, Optional
import logging
from sentence_transformers import SentenceTransformer
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import pickle
import hashlib


try:
    nltk.data.find('tokenizers/punkt')
    nltk.data.find('corpora/stopwords')
    nltk.data.find('corpora/wordnet')
except LookupError:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('wordnet', quiet=True)


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


st.set_page_config(
    page_title="AI Document Q&A Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)


st.markdown("""
<style>
    .main-header {
        text-align: center;
        color: #1f77b4;
        margin-bottom: 30px;
        font-size: 2.5rem;
        font-weight: bold;
    }
    
    .chat-container {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        border-radius: 15px;
        padding: 20px;
        margin: 10px 0;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .user-message {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 15px;
        border-radius: 15px 15px 5px 15px;
        margin: 10px 0;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.2);
    }
    
    .bot-message {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        color: white;
        padding: 15px;
        border-radius: 15px 15px 15px 5px;
        margin: 10px 0;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.2);
    }
    
    .document-info {
        background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
        color: white;
        padding: 15px;
        border-radius: 10px;
        margin: 10px 0;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    
    .similarity-chunk {
        background: #f8f9fa;
        border: 1px solid #e9ecef;
        border-left: 4px solid #17a2b8;
        padding: 10px;
        margin: 5px 0;
        border-radius: 5px;
        font-size: 0.9em;
    }
    
    .confidence-score {
        background: linear-gradient(90deg, #11998e, #38ef7d);
        color: white;
        padding: 5px 10px;
        border-radius: 15px;
        font-size: 0.8em;
        font-weight: bold;
        display: inline-block;
        margin: 5px 0;
    }
    
    .no-answer {
        background: #fff3cd;
        border: 1px solid #ffeaa7;
        color: #856404;
        padding: 15px;
        border-radius: 10px;
        margin: 10px 0;
    }
    
    .processing {
        background: #e3f2fd;
        border-left: 4px solid #2196f3;
        padding: 10px;
        margin: 10px 0;
        border-radius: 5px;
    }
</style>
""", unsafe_allow_html=True)

class AdvancedDocumentProcessor:
    """Enhanced document processing with better text extraction"""
    
    @staticmethod
    def extract_text_from_pdf(file) -> str:
        """Extract text from PDF with better handling"""
        try:
            file.seek(0)
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # Clean and normalize text
                        page_text = re.sub(r'\s+', ' ', page_text)
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file) -> str:
        """Extract text from DOCX with better handling"""
        try:
            file.seek(0)
            doc = docx.Document(file)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file) -> str:
        """Extract text from TXT file"""
        try:
            file.seek(0)
            content = file.read()
            if isinstance(content, bytes):
                # Try different encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        return content.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                return content.decode('utf-8', errors='ignore')
            return content
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            raise Exception(f"Failed to extract text from TXT: {str(e)}")

class IntelligentTextChunker:
    """Advanced text chunking with semantic awareness"""
    
    def __init__(self):
        self.lemmatizer = WordNetLemmatizer()
        self.stop_words = set(stopwords.words('english'))
    
    def preprocess_text(self, text: str) -> str:
        """Clean and preprocess text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep sentence structure
        text = re.sub(r'[^\w\s\.\!\?\,\;\:\-\(\)]', ' ', text)
        
        # Fix common OCR errors
        text = re.sub(r'\b\d+\s*\.\s*\d+\b', lambda m: m.group().replace(' ', ''), text)
        
        return text.strip()
    
    def chunk_text_semantically(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict]:
        """Create semantic chunks with metadata"""
        if not text or len(text.strip()) < 50:
            return []
        
        # Preprocess text
        text = self.preprocess_text(text)
        
        # Split into sentences
        sentences = sent_tokenize(text)
        
        chunks = []
        current_chunk = ""
        current_sentences = []
        
        for i, sentence in enumerate(sentences):
            # Calculate if adding this sentence exceeds chunk size
            potential_chunk = current_chunk + " " + sentence if current_chunk else sentence
            
            if len(potential_chunk) > chunk_size and current_chunk:
                # Save current chunk
                chunk_info = {
                    'text': current_chunk.strip(),
                    'sentences': current_sentences.copy(),
                    'start_sentence': current_sentences[0] if current_sentences else i,
                    'end_sentence': current_sentences[-1] if current_sentences else i,
                    'word_count': len(current_chunk.split()),
                    'char_count': len(current_chunk)
                }
                chunks.append(chunk_info)
                
                # Start new chunk with overlap
                if len(current_sentences) > 1:
                    overlap_sentences = current_sentences[-2:] if len(current_sentences) > 2 else current_sentences[-1:]
                    current_chunk = " ".join([sentences[idx] for idx in overlap_sentences]) + " " + sentence
                    current_sentences = overlap_sentences + [i]
                else:
                    current_chunk = sentence
                    current_sentences = [i]
            else:
                current_chunk = potential_chunk
                current_sentences.append(i)
        
        # Add the last chunk
        if current_chunk:
            chunk_info = {
                'text': current_chunk.strip(),
                'sentences': current_sentences,
                'start_sentence': current_sentences[0] if current_sentences else len(sentences)-1,
                'end_sentence': current_sentences[-1] if current_sentences else len(sentences)-1,
                'word_count': len(current_chunk.split()),
                'char_count': len(current_chunk)
            }
            chunks.append(chunk_info)
        
        # Filter out very short chunks
        chunks = [chunk for chunk in chunks if chunk['word_count'] > 10]
        
        return chunks

class SmartQASystem:
    """Intelligent Q&A system with multiple retrieval strategies"""
    
    def __init__(self):
        self.tfidf_vectorizer = TfidfVectorizer(
            max_features=10000,
            stop_words='english',
            ngram_range=(1, 3),
            min_df=1,
            max_df=0.95,
            sublinear_tf=True
        )
        
        # Initialize sentence transformer for semantic similarity
        try:
            self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
        except Exception as e:
            logger.warning(f"Could not load sentence transformer: {e}")
            self.sentence_model = None
        
        self.document_chunks = []
        self.tfidf_matrix = None
        self.semantic_embeddings = None
        self.document_metadata = {}
        self.chat_history = []
        
        # Question type patterns
        self.question_patterns = {
            'what': r'\b(what|which)\b',
            'how': r'\b(how)\b',
            'when': r'\b(when)\b',
            'where': r'\b(where)\b',
            'why': r'\b(why)\b',
            'who': r'\b(who)\b',
            'definition': r'\b(define|definition|meaning|means)\b',
            'list': r'\b(list|enumerate|mention|name)\b',
            'comparison': r'\b(compare|difference|similar|versus|vs)\b',
            'explanation': r'\b(explain|describe|elaborate|detail)\b'
        }
    
    def load_document(self, file) -> bool:
        """Load and process document with advanced chunking"""
        try:
            # Extract text based on file type
            filename = file.name
            file_extension = filename.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                text = AdvancedDocumentProcessor.extract_text_from_pdf(file)
            elif file_extension == 'docx':
                text = AdvancedDocumentProcessor.extract_text_from_docx(file)
            elif file_extension == 'txt':
                text = AdvancedDocumentProcessor.extract_text_from_txt(file)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            if not text or len(text.strip()) < 100:
                raise ValueError("Document appears to be empty or too short")
            
            # Create intelligent chunks
            chunker = IntelligentTextChunker()
            self.document_chunks = chunker.chunk_text_semantically(text, chunk_size=800, overlap=150)
            
            if not self.document_chunks:
                raise ValueError("Could not create meaningful chunks from document")
            
            # Create TF-IDF matrix
            chunk_texts = [chunk['text'] for chunk in self.document_chunks]
            self.tfidf_matrix = self.tfidf_vectorizer.fit_transform(chunk_texts)
            
            # Create semantic embeddings if available
            if self.sentence_model:
                try:
                    self.semantic_embeddings = self.sentence_model.encode(chunk_texts)
                except Exception as e:
                    logger.warning(f"Could not create semantic embeddings: {e}")
                    self.semantic_embeddings = None
            
            # Store metadata
            self.document_metadata = {
                'filename': filename,
                'file_size': file.size if hasattr(file, 'size') else len(text),
                'upload_time': datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'total_chunks': len(self.document_chunks),
                'total_characters': len(text),
                'total_words': len(text.split()),
                'average_chunk_size': np.mean([chunk['word_count'] for chunk in self.document_chunks])
            }
            
            # Clear previous chat history
            self.chat_history = []
            
            return True
            
        except Exception as e:
            logger.error(f"Error loading document: {e}")
            st.error(f"Error loading document: {str(e)}")
            return False
    
    def classify_question_type(self, question: str) -> str:
        """Classify the type of question being asked"""
        question_lower = question.lower()
        
        for q_type, pattern in self.question_patterns.items():
            if re.search(pattern, question_lower):
                return q_type
        
        return 'general'
    
    def find_relevant_chunks_hybrid(self, query: str, top_k: int = 5) -> List[Tuple[Dict, float, str]]:
        """Find relevant chunks using hybrid approach (TF-IDF + Semantic)"""
        if not self.document_chunks:
            return []
        
        try:
            # TF-IDF similarity
            query_tfidf = self.tfidf_vectorizer.transform([query])
            tfidf_similarities = cosine_similarity(query_tfidf, self.tfidf_matrix).flatten()
            
            # Semantic similarity (if available)
            semantic_similarities = None
            if self.sentence_model and self.semantic_embeddings is not None:
                try:
                    query_embedding = self.sentence_model.encode([query])
                    semantic_similarities = cosine_similarity(query_embedding, self.semantic_embeddings).flatten()
                except Exception as e:
                    logger.warning(f"Error computing semantic similarity: {e}")
            
            # Combine similarities
            combined_scores = []
            for i, chunk in enumerate(self.document_chunks):
                tfidf_score = tfidf_similarities[i]
                semantic_score = semantic_similarities[i] if semantic_similarities is not None else 0
                
                # Weight combination (favor semantic if available)
                if semantic_similarities is not None:
                    combined_score = 0.4 * tfidf_score + 0.6 * semantic_score
                    method = "hybrid"
                else:
                    combined_score = tfidf_score
                    method = "tfidf"
                
                combined_scores.append((chunk, combined_score, method))
            
            # Sort by combined score and return top k
            combined_scores.sort(key=lambda x: x[1], reverse=True)
            
            # Filter by minimum threshold
            relevant_chunks = [(chunk, score, method) for chunk, score, method in combined_scores[:top_k] if score > 0.1]
            
            return relevant_chunks
            
        except Exception as e:
            logger.error(f"Error finding relevant chunks: {e}")
            return []
    
    def extract_key_information(self, query: str, chunks: List[Dict]) -> List[str]:
        """Extract key information from chunks based on query"""
        query_words = set(query.lower().split())
        query_type = self.classify_question_type(query)
        
        key_info = []
        
        for chunk in chunks:
            text = chunk['text']
            sentences = sent_tokenize(text)
            
            # Score sentences based on relevance
            sentence_scores = []
            for sentence in sentences:
                sentence_words = set(sentence.lower().split())
                word_overlap = len(query_words.intersection(sentence_words))
                
                # Boost score based on question type
                type_boost = 1.0
                if query_type == 'definition' and any(word in sentence.lower() for word in ['is', 'are', 'means', 'defined']):
                    type_boost = 1.5
                elif query_type == 'list' and any(word in sentence.lower() for word in ['include', 'such as', 'example']):
                    type_boost = 1.3
                elif query_type == 'how' and any(word in sentence.lower() for word in ['process', 'method', 'way', 'step']):
                    type_boost = 1.4
                
                score = (word_overlap / len(query_words)) * type_boost if query_words else 0
                sentence_scores.append((sentence, score))
            
            # Get top sentences
            sentence_scores.sort(key=lambda x: x[1], reverse=True)
            for sentence, score in sentence_scores[:3]:
                if score > 0.1 and len(sentence) > 20:
                    key_info.append(sentence)
        
        return key_info[:5]  # Return top 5 key sentences
    
    def generate_comprehensive_answer(self, query: str, relevant_chunks: List[Tuple[Dict, float, str]]) -> Dict:
        """Generate a comprehensive answer with confidence scoring"""
        if not relevant_chunks:
            return {
                'answer': "I couldn't find relevant information in the document to answer your question. Please try rephrasing your question or ensure the document contains the information you're looking for.",
                'confidence': 0.0,
                'sources': [],
                'method': 'none'
            }
        
        # Extract key information
        chunks_data = [chunk for chunk, _, _ in relevant_chunks]
        key_sentences = self.extract_key_information(query, chunks_data)
        
        # Calculate overall confidence
        avg_similarity = np.mean([score for _, score, _ in relevant_chunks])
        confidence = min(avg_similarity * 100, 95)  # Cap at 95%
        
        # Determine method used
        methods = [method for _, _, method in relevant_chunks]
        primary_method = max(set(methods), key=methods.count)
        
        # Generate answer
        if key_sentences:
            # Create a coherent answer from key sentences
            answer_parts = []
            seen_content = set()
            
            for sentence in key_sentences:
                # Avoid duplicate information
                sentence_words = set(sentence.lower().split())
                if not any(len(sentence_words.intersection(seen)) > len(sentence_words) * 0.7 for seen in seen_content):
                    answer_parts.append(sentence.strip())
                    seen_content.add(sentence_words)
            
            if answer_parts:
                answer = " ".join(answer_parts)
                
                # Add context if answer is too short
                if len(answer) < 100 and relevant_chunks:
                    context = relevant_chunks[0][0]['text'][:300]
                    answer += f"\n\nAdditional context: {context}..."
            else:
                # Fallback to first chunk
                answer = relevant_chunks[0][0]['text'][:400] + "..."
        else:
            # Fallback to showing relevant content
            answer = relevant_chunks[0][0]['text'][:400] + "..."
        
        # Prepare source information
        sources = []
        for chunk, score, method in relevant_chunks[:3]:
            sources.append({
                'text': chunk['text'][:200] + "..." if len(chunk['text']) > 200 else chunk['text'],
                'similarity': score,
                'method': method,
                'word_count': chunk['word_count']
            })
        
        return {
            'answer': answer,
            'confidence': confidence,
            'sources': sources,
            'method': primary_method,
            'question_type': self.classify_question_type(query)
        }
    
    def chat(self, user_input: str) -> Dict:
        """Main chat function with comprehensive response"""
        if not self.document_chunks:
            return {
                'answer': "Please upload and process a document first.",
                'confidence': 0.0,
                'sources': [],
                'method': 'none'
            }
        
        # Find relevant chunks
        relevant_chunks = self.find_relevant_chunks_hybrid(user_input, top_k=7)
        
        # Generate comprehensive answer
        response = self.generate_comprehensive_answer(user_input, relevant_chunks)
        
        # Add to chat history
        chat_entry = {
            'user_query': user_input,
            'bot_response': response['answer'],
            'confidence': response['confidence'],
            'sources': response['sources'],
            'method': response['method'],
            'question_type': response['question_type'],
            'timestamp': datetime.datetime.now().strftime("%H:%M:%S")
        }
        
        self.chat_history.append(chat_entry)
        
        return response

def main():
    """Main Streamlit application"""
    st.markdown("<h1 class='main-header'>🤖 AI Document Q&A Assistant</h1>", unsafe_allow_html=True)
    
    # Initialize session state
    if 'qa_system' not in st.session_state:
        st.session_state.qa_system = SmartQASystem()
    
    if 'processing' not in st.session_state:
        st.session_state.processing = False
    
    # Sidebar for document upload and settings
    with st.sidebar:
        st.header("📄 Document Management")
        
        uploaded_file = st.file_uploader(
            "Upload Document",
            type=['pdf', 'docx', 'txt'],
            help="Supported formats: PDF, DOCX, TXT (Max 200MB)"
        )
        
        if uploaded_file is not None:
            if st.button("🚀 Process Document", type="primary"):
                st.session_state.processing = True
                
                with st.spinner("🔄 Processing document... This may take a few moments"):
                    progress_bar = st.progress(0)
                    
                    # Simulate processing steps
                    progress_bar.progress(20)
                    st.info("📖 Extracting text from document...")
                    
                    progress_bar.progress(50)
                    st.info("🔍 Creating intelligent chunks...")
                    
                    progress_bar.progress(80)
                    st.info("🧠 Building AI index...")
                    
                    success = st.session_state.qa_system.load_document(uploaded_file)
                    progress_bar.progress(100)
                    
                    if success:
                        st.success("✅ Document processed successfully!")
                        st.balloons()
                    else:
                        st.error("❌ Failed to process document")
                
                st.session_state.processing = False
                st.rerun()
        
        # Document info
        if st.session_state.qa_system.document_metadata:
            st.header("📊 Document Analytics")
            metadata = st.session_state.qa_system.document_metadata
            
            st.markdown(f"""
            <div class='document-info'>
                <strong>📁 File:</strong> {metadata['filename']}<br>
                <strong>⏰ Processed:</strong> {metadata['upload_time']}<br>
                <strong>📄 Chunks:</strong> {metadata['total_chunks']}<br>
                <strong>📝 Words:</strong> {metadata['total_words']:,}<br>
                <strong>🔤 Characters:</strong> {metadata['total_characters']:,}<br>
                <strong>📊 Avg Chunk Size:</strong> {metadata['average_chunk_size']:.0f} words
            </div>
            """, unsafe_allow_html=True)
            
            # Embedding model info
            st.header("🤖 AI Model Info")
            model_info = "🔍 **Search:** TF-IDF + Semantic<br>🧠 **Embeddings:** all-MiniLM-L6-v2<br>📊 **Similarity:** Cosine"
            st.markdown(f'<div class="document-info">{model_info}</div>', unsafe_allow_html=True)
        
        # Settings
        st.header("⚙️ Settings")
        show_sources = st.checkbox("📚 Show Source References", value=True)
        show_confidence = st.checkbox("📊 Show Confidence Scores", value=True)
        detailed_mode = st.checkbox("🔍 Detailed Analysis Mode", value=False)
        
        if st.button("🗑️ Clear Chat History"):
            st.session_state.qa_system.chat_history = []
            st.success("Chat history cleared!")
            st.rerun()
    
    # Main chat interface
    if not st.session_state.qa_system.document_chunks:
        st.markdown("""
        <div class='chat-container'>
            <h2>🎯 Welcome to AI Document Q&A Assistant</h2>
            <p>Your intelligent document companion powered by advanced AI!</p>
            
            <h3>🚀 How to get started:</h3>
            <ol>
                <li>📤 Upload your document (PDF, DOCX, or TXT)</li>
                <li>🔄 Click "Process Document" to analyze content</li>
                <li>💬 Ask questions about your document</li>
                <li>🎯 Get accurate answers with confidence scores</li>
            </ol>
            
            <h3>✨ Features:</h3>
            <ul>
                <li>🧠 <strong>Hybrid AI Search:</strong> TF-IDF + Semantic Understanding</li>
                <li>📊 <strong>Confidence Scoring:</strong> Know how reliable each answer is</li>
                <li>📚 <strong>Source References:</strong> See exactly where answers come from</li>
                <li>🔍 <strong>Smart Chunking:</strong> Intelligent text segmentation</li>
                <li>💡 <strong>Question Classification:</strong> Optimized for different question types</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        return
    
    # Chat interface
    st.header("💬 Ask Your Questions")
    
    # Input form
    with st.form(key='question_form'):
        user_question = st.text_area(
            "What would you like to know about your document?",
            placeholder="Examples:\n• What is the main topic of this document?\n• Can you summarize the key points?\n• How does the author explain [specific concept]?",
            height=100
        )
        
        col1, col2 = st.columns([3, 1])
        with col1:
            submitted = st.form_submit_button("🔍 Get Answer", type="primary")
        with col2:
            if st.form_submit_button("🎲 Random Question"):
                sample_questions = [
                    "What is the main topic of this document?",
                    "Can you summarize the key points?",
                    "What are the most important findings?",
                    "Who are the main people mentioned?",
                    "What conclusions does the author draw?"
                ]
                user_question = np.random.choice(sample_questions)
                submitted = True
        
        if submitted and user_question.strip():
            with st.spinner("🤖 AI is analyzing your question..."):
                response = st.session_state.qa_system.chat(user_question.strip())
                
                # Display answer with styling
                confidence = response['confidence']
                confidence_color = "green" if confidence > 70 else "orange" if confidence > 40 else "red"
                
                st.markdown(f"""
                <div class='bot-message'>
                    <strong>🤖 AI Assistant:</strong><br>
                    {response['answer']}
                </div>
                """, unsafe_allow_html=True)
                
                if show_confidence:
                    st.markdown(f"""
                    <div class='confidence-score' style='background: {confidence_color}'>
                        📊 Confidence: {confidence:.1f}% | Method: {response['method'].upper()} | Type: {response['question_type'].upper()}
                    </div>
                    """, unsafe_allow_html=True)
                
                if show_sources and response['sources']:
                    st.markdown("**📚 Source References:**")
                    for i, source in enumerate(response['sources'], 1):
                        st.markdown(f"""
                        <div class='similarity-chunk'>
                            <strong>📄 Source {i} (Similarity: {source['similarity']:.2f})</strong><br>
                            {source['text']}
                        </div>
                        """, unsafe_allow_html=True)
            
            st.rerun()
    
    # Chat history
    if st.session_state.qa_system.chat_history:
        st.header("📝 Chat History")
        
        for i, chat in enumerate(reversed(st.session_state.qa_system.chat_history)):
            with st.expander(f"💬 Q{len(st.session_state.qa_system.chat_history)-i}: {chat['user_query'][:60]}... | {chat['timestamp']}"):
                # User question